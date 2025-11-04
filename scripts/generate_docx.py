/* LABELED_BY_TOOL
 * File: scripts/generate_docx.py
 * Inferred role: Project file — please open to see specific role
 * Note: auto-generated label. Please edit the file for a more accurate description. */

"""
Simple Markdown -> DOCX converter for the FDD professional report.
Usage:
  python -m pip install python-docx markdown
  python scripts/generate_docx.py ../reports/Professional_FDD_Report.md ../reports/Professional_FDD_Report.docx

This script performs a minimal mapping of Markdown headings to Word heading styles
and preserves paragraphs and simple tables. It is intentionally lightweight to avoid
complex dependencies.
"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("python-docx is required. Install with: python -m pip install python-docx")
    sys.exit(1)


def add_table_from_markdown(doc, lines, start_idx):
    """Parse a simple pipe-separated markdown table block starting at start_idx.
    Returns the index after the table.
    """
    # Collect table lines
    tbl_lines = []
    i = start_idx
    while i < len(lines) and '|' in lines[i]:
        line = lines[i].strip()
        if line:
            tbl_lines.append(line)
        i += 1
    if not tbl_lines:
        return start_idx
    # Split rows
    rows = [ [c.strip() for c in r.strip('|').split('|')] for r in tbl_lines ]
    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            table.cell(r_idx, c_idx).text = cell
    return i


def md_to_docx(md_path: Path, out_path: Path):
    lines = md_path.read_text(encoding='utf-8').splitlines()
    doc = Document()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith('---'):
            # horizontal rule -> small break
            doc.add_paragraph('')
            i += 1
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if '|' in line and line.strip().startswith('|'):
            i = add_table_from_markdown(doc, lines, i)
            continue
        if line.strip() == '':
            # paragraph break
            doc.add_paragraph('')
            i += 1
            continue
        # Paragraph or list
        # handle simple bullet lines
        if line.strip().startswith('- '):
            p = doc.add_paragraph(line.strip()[2:].strip(), style='List Bullet')
            i += 1
            continue
        if line.strip().startswith('1.'):
            p = doc.add_paragraph(line.strip()[3:].strip(), style='List Number')
            i += 1
            continue
        # otherwise simple paragraph: accumulate until blank line
        para_lines = [line]
        j = i + 1
        while j < len(lines) and lines[j].strip() != '':
            para_lines.append(lines[j])
            j += 1
        text = '\n'.join(para_lines).strip()
        doc.add_paragraph(text)
        i = j
    # Save
    doc.save(str(out_path))
    print(f"Wrote DOCX to {out_path}")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python scripts/generate_docx.py <input.md> <output.docx>")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])
    if not md_path.exists():
        print(f"Input file not found: {md_path}")
        sys.exit(1)
    md_to_docx(md_path, out_path)
